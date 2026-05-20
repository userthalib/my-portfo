import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = docx.Document()
    
    # Set document margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Color Palette Constants (Premium Modern Theme: Indigo/Purple/Slate)
    COLOR_PRIMARY = RGBColor(99, 102, 241)     # Indigo (#6366F1)
    COLOR_SECONDARY = RGBColor(6, 182, 212)   # Cyan (#06B6D4)
    COLOR_TEXT = RGBColor(31, 41, 55)         # Charcoal (#1F2937)
    COLOR_MUTED = RGBColor(100, 116, 139)     # Slate Muted (#64748B)

    # XML helpers for cell shading and left-hand borders
    def set_cell_background(cell, fill_hex):
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shd)

    def set_cell_left_border(cell, border_color_hex, size_eighths=12):
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
            f'<w:left w:val="single" w:sz="{size_eighths}" w:space="0" w:color="{border_color_hex}"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
            f'<w:right w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
            f'</w:tcBorders>'
        )
        cell._tc.get_or_add_tcPr().append(borders)

    # Custom Helper for paragraph addition with exact formatting control
    def add_p(text="", font_name="Calibri", font_size=11, bold=False, italic=False, color=COLOR_TEXT, before=0, after=6, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.15):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = line_spacing
        
        if text:
            run = p.add_run(text)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = bold
            run.italic = italic
            run.font.color.rgb = color
            
        return p

    def add_run(p, text, bold=False, italic=False, color=COLOR_TEXT, font_size=11, font_name="Calibri"):
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = color
        return run

    # Helper to add stylized Callout block
    def add_callout(text, title="NOTE", fill_hex="EEF2F6", border_hex="6366F1"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.autofit = False
        tbl.columns[0].width = Inches(6.5)
        cell = tbl.cell(0, 0)
        
        set_cell_background(cell, fill_hex)
        set_cell_left_border(cell, border_hex, size_eighths=16)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        
        run_title = p.add_run(f"[{title}] ")
        run_title.font.name = "Calibri"
        run_title.font.size = Pt(10)
        run_title.bold = True
        run_title.font.color.rgb = COLOR_PRIMARY
        
        run_text = p.add_run(text)
        run_text.font.name = "Calibri"
        run_text.font.size = Pt(10)
        run_text.italic = True
        run_text.font.color.rgb = COLOR_TEXT
        
        # Add space after table
        add_p(before=0, after=8)

    # Helper to add Code block (shaded background, Consolas font, left border)
    def add_code(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.autofit = False
        tbl.columns[0].width = Inches(6.5)
        cell = tbl.cell(0, 0)
        
        set_cell_background(cell, "F8FAFC")
        set_cell_left_border(cell, "64748B", size_eighths=8)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.05
        
        run = p.add_run(code_text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(15, 23, 42)
        
        # Add space after table
        add_p(before=0, after=8)

    # ─────────────────────────────────────────────────────────────
    # COVER PAGE (Modern Minimalist Style)
    # ─────────────────────────────────────────────────────────────
    
    # Vertical spacing to center visual content roughly
    for _ in range(3):
        add_p(after=12)
        
    p_title = add_p("PORTFOLIO WEBSITE", font_name="Calibri", font_size=28, bold=True, color=COLOR_PRIMARY, align=WD_ALIGN_PARAGRAPH.CENTER)
    p_title_sub = add_p("User Guide & Customization Manual", font_name="Calibri", font_size=20, bold=False, color=COLOR_SECONDARY, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=12)
    
    p_desc = add_p(
        "A comprehensive non-developer and developer guide to managing, personalizing, configuring, and deploying your high-performance portfolio website.",
        font_size=12, italic=True, color=COLOR_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=24
    )
    p_desc.paragraph_format.left_indent = Inches(1)
    p_desc.paragraph_format.right_indent = Inches(1)
    
    for _ in range(8):
        add_p(after=12)
        
    add_p("Prepared For: Irfan Thalib", font_size=11, bold=True, color=COLOR_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p("Technical Infrastructure Document", font_size=10, color=COLOR_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p("Last Updated: May 2026", font_size=10, color=COLOR_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    
    # Page Break after Cover
    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # SECTION 1: INTRODUCTION
    # ─────────────────────────────────────────────────────────────
    h1 = add_p("1. Introduction & Overview", font_size=18, bold=True, color=COLOR_PRIMARY, before=18, after=12)
    h1.paragraph_format.keep_with_next = True
    
    p1 = add_p(before=0, after=8)
    add_run(p1, "Welcome to your new portfolio project! This application has been designed as a premium, state-of-the-art landing page. Utilizing modern design aesthetics like ambient glow effects, modular grids, responsive timeline cards, and interactive hover highlights, the website creates a professional first impression.")
    
    p2 = add_p(before=0, after=8)
    add_run(p2, "The project is built on ")
    add_run(p2, "Next.js (App Router)", bold=True)
    add_run(p2, ", styling is driven by ")
    add_run(p2, "Tailwind CSS v4", bold=True)
    add_run(p2, ", and interactive components are powered by ")
    add_run(p2, "Framer Motion", bold=True)
    add_run(p2, " (including scroll progress tracking and custom spotlight cursor effects). The backend contact forms link directly into ")
    add_run(p2, "Formspree", bold=True)
    add_run(p2, " for simple, zero-setup serverless message delivery.")

    add_callout(
        "For ease of maintenance, all components have been thoroughly documented directly inside the code using JSDoc. This manual will guide you on how to change names, update descriptions, add project entries, customize styling, and deploy the finished page to production.",
        title="GOAL"
    )

    # ─────────────────────────────────────────────────────────────
    # SECTION 2: CODEBASE ARCHITECTURE
    # ─────────────────────────────────────────────────────────────
    h2 = add_p("2. Codebase Structure & Architecture", font_size=18, bold=True, color=COLOR_PRIMARY, before=18, after=12)
    h2.paragraph_format.keep_with_next = True
    
    p_arch = add_p(before=0, after=8)
    add_run(p_arch, "The portfolio utilizes a Next.js App Router structure. Below is a map of the important directories and files:")
    
    # Add bulleted lists of structure
    def add_bullet(bold_term, definition, path_text=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        
        run_term = p.add_run(bold_term)
        run_term.bold = True
        run_term.font.color.rgb = COLOR_TEXT
        
        if path_text:
            run_path = p.add_run(f" ({path_text})")
            run_path.font.color.rgb = COLOR_MUTED
            run_path.font.name = "Consolas"
            run_path.font.size = Pt(9.5)
            
        run_def = p.add_run(f" — {definition}")
        run_def.font.color.rgb = COLOR_TEXT
        
    add_bullet("Page Layout Wrapper", "Sets fonts (Geist & Space Grotesk), layout bounds, and visual background radial lights.", "app/layout.tsx")
    add_bullet("Home Page Entry", "Assembles page blocks under a unified Section layout wrapper.", "app/page.tsx")
    add_bullet("Navigation Bar", "Handles scroll tracking, mobile responsive drawer toggle, and animated active indicator.", "app/components/Navbar.tsx")
    add_bullet("Hero Section", "Landing area displaying titles, custom profile picture, social links, and typewriter animation loop.", "app/components/Hero.tsx")
    add_bullet("About Section", "Highlights developer bio, quantifications grid, and dynamically colored skill badges.", "app/components/About.tsx")
    add_bullet("Services Grid", "Grid displaying services, custom graphics, and specialized technology badges.", "app/components/Services.tsx")
    add_bullet("Journey Timeline", "Alternating timeline detailing career events, education milestones, and tech stacks.", "app/components/Experience.tsx")
    add_bullet("Projects Bento Grid", "Renders project grids, external links, and fallback brackets.", "app/components/Projects.tsx")
    add_bullet("Contact & Form", "Includes email form linked directly to the Formspree endpoint.", "app/components/Contact.tsx")
    add_bullet("Footer Block", "Provides footer anchors, back-to-top button, and copyright labels.", "app/components/Footer.tsx")
    add_bullet("Scroll Progress Bar", "Framer Motion indicator showing scroll status at top of the window.", "app/components/ProgressBar.tsx")
    add_bullet("Spotlight Card Wrapper", "Calculates local mouse coordinate tracking to display hover radial gradient lighting.", "app/components/SpotlightCard.tsx")

    # ─────────────────────────────────────────────────────────────
    # SECTION 3: STEP-BY-STEP CUSTOMIZATION
    # ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    h3 = add_p("3. Step-by-Step Customization", font_size=18, bold=True, color=COLOR_PRIMARY, before=18, after=12)
    h3.paragraph_format.keep_with_next = True
    
    # 3.1 Personal Info & Typewriter
    h3_1 = add_p("3.1 Name, Bio, and Typewriter Roles", font_size=14, bold=True, color=COLOR_SECONDARY, before=12, after=6)
    h3_1.paragraph_format.keep_with_next = True
    
    p_3_1 = add_p(before=0, after=8)
    add_run(p_3_1, "All primary credentials are configured as standard constants in the component files. To edit the roles cycled in the typewriter header, open ")
    add_run(p_3_1, "app/components/Hero.tsx", bold=True, font_name="Consolas", font_size=9.5, color=COLOR_MUTED)
    add_run(p_3_1, " and locate the ")
    add_run(p_3_1, "ROLES", bold=True, font_name="Consolas", font_size=9.5)
    add_run(p_3_1, " constant:")
    
    add_code(
        'const ROLES = [\n'
        '  "Full Stack Developer",\n'
        '  "React Engineer",\n'
        '  "Laravel Developer",\n'
        '  "UI/UX Enthusiast",\n'
        '];'
    )
    
    p_3_1_b = add_p(before=0, after=8)
    add_run(p_3_1_b, "You can add, edit, or remove roles as text items in this array. The typewriter engine will automatically cycle through all available items.")
    
    # 3.2 Profile Photo
    h3_2 = add_p("3.2 Setting Your Profile Photo", font_size=14, bold=True, color=COLOR_SECONDARY, before=12, after=6)
    h3_2.paragraph_format.keep_with_next = True
    
    p_3_2 = add_p(before=0, after=8)
    add_run(p_3_2, "By default, the website renders a modern initials emblem (")
    add_run(p_3_2, "IT", bold=True)
    add_run(p_3_2, "). To replace this initials block with your photo:")
    
    add_bullet("Save Image", "Add your portrait photo directly into the public directory of the project, e.g., as public/profile.jpg.")
    add_bullet("Configure Path", "Open app/components/Hero.tsx and locate PHOTO_SRC. Change its value to the relative filename path, e.g.:", "app/components/Hero.tsx")
    add_code('const PHOTO_SRC: string | null = "/profile.jpg";')
    
    add_callout(
        "If you want to pull your image from an external CDN host (e.g. https://myhost.com/profile.jpg), specify the absolute URL in PHOTO_SRC and configure the domain in your next.config.ts images section so Next.js Image loader allows remote downloads.",
        title="REMOTE DOMAINS"
    )

    # 3.3 Skill Tags & Custom Colors
    h3_3 = add_p("3.3 Stats, Tech Stack, & Custom Badges", font_size=14, bold=True, color=COLOR_SECONDARY, before=12, after=6)
    h3_3.paragraph_format.keep_with_next = True
    
    p_3_3 = add_p(before=0, after=8)
    add_run(p_3_3, "In ")
    add_run(p_3_3, "app/components/About.tsx", bold=True, font_name="Consolas", font_size=9.5, color=COLOR_MUTED)
    add_run(p_3_3, ", you can configure metrics, tech tags, and brand colors. Open the file and adjust:")
    
    add_code(
        '// Quantitative experience markers\n'
        'const STATS = [\n'
        '  { num: "3+",  label: "Years Experience", color: "text-purple-400" },\n'
        '  { num: "15+", label: "Projects Shipped",  color: "text-cyan-400"   },\n'
        '  { num: "5+",  label: "Happy Clients",     color: "text-pink-400"   },\n'
        '  { num: "∞",   label: "Cups of Coffee",    color: "text-amber-400"  },\n'
        '];'
    )
    
    p_3_3_b = add_p(before=0, after=8)
    add_run(p_3_3_b, "To ensure technology tags look visually distinct, modify the ")
    add_run(p_3_3_b, "TECH_COLORS", bold=True, font_name="Consolas", font_size=9.5)
    add_run(p_3_3_b, " mapping to apply specific Tailwind classes to specific tags:")
    
    add_code(
        'const TECH_COLORS: Record<string, string> = {\n'
        '  JavaScript:  "border-yellow-500/20 bg-yellow-500/8  text-yellow-300",\n'
        '  TypeScript:  "border-blue-500/20   bg-blue-500/8    text-blue-300",\n'
        '  React:       "border-cyan-500/20   bg-cyan-500/8    text-cyan-300",\n'
        '  Laravel:     "border-red-500/20    bg-red-500/8     text-red-300",\n'
        '  // Add or modify mappings...\n'
        '};'
    )

    # 3.4 Services
    h3_4 = add_p("3.4 Specialization Services", font_size=14, bold=True, color=COLOR_SECONDARY, before=12, after=6)
    h3_4.paragraph_format.keep_with_next = True
    
    p_3_4 = add_p(before=0, after=8)
    add_run(p_3_4, "To update services, go to ")
    add_run(p_3_4, "app/components/Services.tsx", bold=True, font_name="Consolas", font_size=9.5, color=COLOR_MUTED)
    add_run(p_3_4, " and edit the ")
    add_run(p_3_4, "SERVICES", bold=True, font_name="Consolas", font_size=9.5)
    add_run(p_3_4, " list. Each service configuration contains an icon reference (from Lucide React), a color profile, and metadata tags:")
    
    add_code(
        '{\n'
        '  icon: Globe,\n'
        '  title: "Web Application Development",\n'
        '  description: "End-to-end web apps built with React & Next.js on the front, Laravel or Node on the back.",\n'
        '  tags: ["React", "Next.js", "Laravel"],\n'
        '  gradient: "from-purple-500/20 to-violet-500/20",\n'
        '  iconColor: "text-purple-400",\n'
        '  iconBg: "bg-purple-500/10",\n'
        '}'
    )

    # 3.5 Experience & Education
    h3_5 = add_p("3.5 Experience & Education Timeline", font_size=14, bold=True, color=COLOR_SECONDARY, before=12, after=6)
    h3_5.paragraph_format.keep_with_next = True
    
    p_3_5 = add_p(before=0, after=8)
    add_run(p_3_5, "Timeline records reside in ")
    add_run(p_3_5, "app/components/Experience.tsx", bold=True, font_name="Consolas", font_size=9.5, color=COLOR_MUTED)
    add_run(p_3_5, ". Modify the ")
    add_run(p_3_5, "TIMELINE", bold=True, font_name="Consolas", font_size=9.5)
    add_run(p_3_5, " array list structure to update companies, courses, and job milestones. Toggle the ")
    add_run(p_3_5, "highlight: true", bold=True, font_name="Consolas", font_size=9.5)
    add_run(p_3_5, " property to apply a purple accent border glow to that specific timeline card:")
    
    add_code(
        '{\n'
        '  type: "work", // or "education"\n'
        '  period: "2024 — Present",\n'
        '  role: "Full Stack Developer",\n'
        '  company: "Freelance",\n'
        '  companyUrl: "https://...", // optional\n'
        '  description: "Building end-to-end web applications...",\n'
        '  tech: ["React", "Next.js", "Laravel", "PostgreSQL"],\n'
        '  highlight: true,\n'
        '}'
    )

    # 3.6 Projects Bento Grid
    h3_6 = add_p("3.6 Projects Bento Grid", font_size=14, bold=True, color=COLOR_SECONDARY, before=12, after=6)
    h3_6.paragraph_format.keep_with_next = True
    
    p_3_6 = add_p(before=0, after=8)
    add_run(p_3_6, "The portfolio uses a Bento grid layout in ")
    add_run(p_3_6, "app/components/Projects.tsx", bold=True, font_name="Consolas", font_size=9.5, color=COLOR_MUTED)
    add_run(p_3_6, ". Featured projects occupy double columns on desktop screens, while standard projects sit in individual blocks side-by-side:")
    
    add_bullet("Featured items", "Set featured: true in the project object. This will span two columns.", "app/components/Projects.tsx")
    add_bullet("Project links", "Set liveUrl, githubUrl, videoUrl, or docsUrl. If set, they will automatically render buttons.", "app/components/Projects.tsx")
    add_bullet("Project Image", "Assign image: \"/projects/my-project.png\" or set image: null to render a fallback bracket vector gradient illustration.", "app/components/Projects.tsx")

    # 3.7 Contact Formspree integration
    h3_7 = add_p("3.7 Contact Form Integration", font_size=14, bold=True, color=COLOR_SECONDARY, before=12, after=6)
    h3_7.paragraph_format.keep_with_next = True
    
    p_3_7 = add_p(before=0, after=8)
    add_run(p_3_7, "The contact form is configured to send messages directly to your email without any backend server. It uses Formspree. To set up your account:")
    
    add_bullet("Create Formspree ID", "Go to formspree.io, register for a free account, and create a new project form endpoint.")
    add_bullet("Update Action URL", "Open app/components/Contact.tsx, locate the form element, and replace the form action endpoint URL with your Formspree form ID URL:", "app/components/Contact.tsx")
    add_code('<form action="https://formspree.io/f/YOUR_FORM_ID" method="POST" className="space-y-5">')

    # ─────────────────────────────────────────────────────────────
    # SECTION 4: LOCAL DEVELOPMENT & DEPLOYMENT
    # ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    h4 = add_p("4. Local Development & Deployment", font_size=18, bold=True, color=COLOR_PRIMARY, before=18, after=12)
    h4.paragraph_format.keep_with_next = True
    
    h4_1 = add_p("4.1 Local Server Environment", font_size=14, bold=True, color=COLOR_SECONDARY, before=12, after=6)
    h4_1.paragraph_format.keep_with_next = True
    
    p_4_1 = add_p(before=0, after=8)
    add_run(p_4_1, "To install project dependencies and start the local development server on your machine, perform the following commands in your shell terminal:")
    
    add_code(
        '# 1. Install all package dependencies\n'
        'npm install\n\n'
        '# 2. Run local development web server\n'
        'npm run dev'
    )
    
    p_4_1_b = add_p(before=0, after=8)
    add_run(p_4_1_b, "Once the development server is active, open ")
    add_run(p_4_1_b, "http://localhost:3000", bold=True, font_name="Consolas", font_size=9.5)
    add_run(p_4_1_b, " in your browser to view your website changes live. The server automatically rebuilds the layout on file updates.")

    # 4.2 Production deployment (Vercel)
    h4_2 = add_p("4.2 Production Deployment (Vercel)", font_size=14, bold=True, color=COLOR_SECONDARY, before=12, after=6)
    h4_2.paragraph_format.keep_with_next = True
    
    p_4_2 = add_p(before=0, after=8)
    add_run(p_4_2, "Deploying your Next.js application to production is easiest using Vercel (the creators of Next.js). Follow this step-by-step procedure:")
    
    add_bullet("Push to Github", "Initialize git, stage all files, commit them, and push your repository to GitHub, GitLab, or Bitbucket.")
    add_bullet("Connect Vercel", "Go to vercel.com, log in, click 'Add New Project', and import your portfolio repository.")
    add_bullet("Configure Build Settings", "Vercel automatically detects Next.js. The default build settings and paths are correct out of the box.")
    add_bullet("Deploy", "Click 'Deploy'. Vercel will compile your code and supply a secure global URL (e.g. yourportfolio.vercel.app).")
    
    add_callout(
        "Next.js build checks will run automatically during Vercel deployments. If you experience TypeScript compile issues, make sure your data constants structure is clean and you have not introduced syntax syntax errors in component code files.",
        title="COMPILE CHECKS"
    )

    # Save to user_documentation.docx
    doc.save("user_documentation.docx")
    print("Documentation docx successfully generated at user_documentation.docx!")

if __name__ == "__main__":
    create_document()
